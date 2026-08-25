# ЧТО ДЕЛАЕТ ЭТОТ СКРИПТ
# Берёт файлы из папки vyhod (это ответы нейросети с метками вроде
# ТЕЛЕФОН-1) и подставляет вместо меток настоящие данные из словаря.
# Результат сохраняет рядом, без приписки -maska в имени файла.
#
# КАК ЗАПУСТИТЬ
# Дважды кликни по этому файлу.
#
# ЧТО ПОЛУЧИТСЯ
# Рядом с каждым файлом в vyhod появится его копия с настоящими данными.
# Если в тексте останутся метки, которых нет в словаре, скрипт предупредит.

import json
import re
from pathlib import Path

SKRIPTY_DIR = Path(__file__).resolve().parent
BASE_DIR = SKRIPTY_DIR.parent
VYHOD_DIR = BASE_DIR / "vyhod"
SLOVAR_PATH = SKRIPTY_DIR / "slovar.json"

TEXTOVYE_RASSHIRENIYA = {".txt", ".md", ".csv", ".json"}

# метки выглядят как СЛОВО-ХВОСТ, например ТЕЛЕФОН-1 или КЛИЕНТ-А
METKA_POHOZHAYA_RE = re.compile(r"\b[А-ЯЁ]{2,}-[А-ЯЁ0-9]+\b")


def load_slovar():
    if not SLOVAR_PATH.exists():
        print(f"Не нашёл файл {SLOVAR_PATH.name} рядом со скриптом.")
        print("Без него нечем подставлять настоящие данные обратно.")
        print("Сначала запусти maska.py на исходных файлах, а уже потом vozvrat.py.")
        raise SystemExit(1)
    try:
        with open(SLOVAR_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
    except (json.JSONDecodeError, UnicodeDecodeError):
        print(f"Не получилось прочитать словарь {SLOVAR_PATH.name} - похоже, файл повреждён.")
        raise SystemExit(1)
    data.setdefault("metka_to_znachenie", {})
    return data


def restore_text(text, metka_to_zn, ne_naideno):
    # заменяем метки, начиная с самых длинных, чтобы не задеть похожие
    for metka in sorted(metka_to_zn.keys(), key=len, reverse=True):
        if metka in text:
            text = text.replace(metka, metka_to_zn[metka])

    for pohozhaya in METKA_POHOZHAYA_RE.findall(text):
        if pohozhaya not in metka_to_zn:
            ne_naideno.add(pohozhaya)

    return text


def restore_paragraph_text(para, novyi_text):
    for run in list(para.runs):
        run.text = ""
    if para.runs:
        para.runs[0].text = novyi_text
    else:
        para.add_run(novyi_text)


def process_docx(path, metka_to_zn, ne_naideno):
    try:
        import docx
    except ImportError:
        print(f"  '{path.name}': чтобы обработать файлы Word (.docx), скрипту нужна ещё одна программа, а её сейчас нет.")
        print(f"  Попроси нейросеть прислать этот ответ обычным текстом и сохрани его как .txt в vyhod - тогда скрипт справится.")
        print(f"  Пока файл пропущен.")
        return None

    try:
        doc = docx.Document(str(path))
    except Exception:
        print(f"  '{path.name}': не получилось открыть файл. Возможно, он повреждён")
        print(f"  или это не настоящий файл Word. Файл пропущен.")
        return None

    def obrabotat_paragrafy(paragraphs):
        for para in paragraphs:
            if para.text.strip():
                novyi = restore_text(para.text, metka_to_zn, ne_naideno)
                if novyi != para.text:
                    restore_paragraph_text(para, novyi)

    obrabotat_paragrafy(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                obrabotat_paragrafy(cell.paragraphs)
    return doc


def process_xlsx(path, metka_to_zn, ne_naideno):
    try:
        import openpyxl
    except ImportError:
        print(f"  '{path.name}': чтобы обработать файлы Excel (.xlsx), скрипту нужна ещё одна программа, а её сейчас нет.")
        print(f"  Попроси нейросеть прислать этот ответ таблицей .csv или обычным текстом - тогда скрипт справится.")
        print(f"  Пока файл пропущен.")
        return None

    try:
        wb = openpyxl.load_workbook(str(path))
    except Exception:
        print(f"  '{path.name}': не получилось открыть файл. Возможно, он повреждён")
        print(f"  или это не настоящий файл Excel. Файл пропущен.")
        return None

    for list_ in wb.worksheets:
        for row in list_.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value.strip():
                    novoe = restore_text(cell.value, metka_to_zn, ne_naideno)
                    if novoe != cell.value:
                        cell.value = novoe
    return wb


def sostavit_imya_rezultata(path):
    stem = path.stem
    if stem.endswith("-maska"):
        novyi_stem = stem[: -len("-maska")]
    else:
        novyi_stem = stem + "-vosstanovlen"
    return f"{novyi_stem}{path.suffix}"


def process_file(path, metka_to_zn, ne_naideno):
    ext = path.suffix.lower()
    out_name = sostavit_imya_rezultata(path)
    out_path = VYHOD_DIR / out_name

    if ext in TEXTOVYE_RASSHIRENIYA:
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            try:
                text = path.read_text(encoding="cp1251")
            except Exception:
                print(f"  '{path.name}': не получилось прочитать текст файла (незнакомая кодировка). Файл пропущен.")
                return False
        novyi_text = restore_text(text, metka_to_zn, ne_naideno)
        out_path.write_text(novyi_text, encoding="utf-8")

    elif ext == ".docx":
        doc = process_docx(path, metka_to_zn, ne_naideno)
        if doc is None:
            return False
        doc.save(str(out_path))

    elif ext == ".xlsx":
        wb = process_xlsx(path, metka_to_zn, ne_naideno)
        if wb is None:
            return False
        wb.save(str(out_path))

    else:
        print(f"  '{path.name}': такой тип файла пока не поддерживается ({ext}).")
        print(f"  Поддерживаются: .txt, .md, .csv, .json, .docx, .xlsx.")
        return False

    return True


def wait_exit():
    try:
        input("\nНажми Enter, чтобы закрыть окно...")
    except Exception:
        pass


def main():
    try:
        VYHOD_DIR.mkdir(exist_ok=True)
        slovar = load_slovar()
        metka_to_zn = slovar.get("metka_to_znachenie", {})

        files = sorted(
            p for p in VYHOD_DIR.iterdir()
            if p.is_file() and not p.name.startswith(".") and not p.stem.endswith("-vosstanovlen")
        )

        if not files:
            print("В папке vyhod нет файлов.")
            print("Сохрани туда ответ нейросети и запусти скрипт снова.")
            wait_exit()
            return

        print(f"Нашёл файлов в vyhod: {len(files)}")
        print("Подставляю настоящие данные...\n")

        ne_naideno = set()
        obrabotano = 0

        for path in files:
            print(f"Обрабатываю: {path.name}")
            try:
                if process_file(path, metka_to_zn, ne_naideno):
                    obrabotano += 1
            except Exception as e:
                print(f"  Не получилось обработать файл '{path.name}': {e}")
                print(f"  Файл пропущен, остальные файлы будут обработаны дальше.")

        print("\nГотово.")
        print(f"Обработано файлов: {obrabotano} из {len(files)}")

        if ne_naideno:
            print("\nВНИМАНИЕ: в тексте остались метки, которых нет в словаре:")
            for metka in sorted(ne_naideno):
                print(f"  {metka}")
            print("Проверь эти места руками - подставь значения сам или пересобери словарь.")
        else:
            print("Меток без пары в словаре не осталось.")

    except Exception as e:
        print(f"Что-то пошло не так: {e}")
        print("Проверь, что рядом со скриптом есть папка vyhod и файл slovar.json, и попробуй ещё раз.")

    wait_exit()


if __name__ == "__main__":
    main()
