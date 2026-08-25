# ЧТО ДЕЛАЕТ ЭТОТ СКРИПТ
# Берёт файлы из папки originaly, прячет в них телефоны, почту, ИНН и другие
# личные данные (заменяет их метками вроде ТЕЛЕФОН-1) и сохраняет
# обезличенные копии в папку vhod. Настоящие данные никогда не покажет.
#
# КАК ЗАПУСТИТЬ
# Дважды кликни по этому файлу.
#
# ЧТО ПОЛУЧИТСЯ
# В папке vhod появятся копии файлов с припиской -maska - их можно
# смело показывать нейросети. Настоящие значения и их метки лягут
# в файл slovar.json рядом с этим скриптом.

import hashlib
import json
import re
from pathlib import Path

# --- где что лежит -----------------------------------------------------

SKRIPTY_DIR = Path(__file__).resolve().parent
BASE_DIR = SKRIPTY_DIR.parent
ORIGINALY_DIR = BASE_DIR / "originaly"
VHOD_DIR = BASE_DIR / "vhod"
SLOVAR_PATH = SKRIPTY_DIR / "slovar.json"

TEXTOVYE_RASSHIRENIYA = {".txt", ".md", ".csv", ".json"}

# --- какие данные ловим шаблонами --------------------------------------

EMAIL_RE = re.compile(r"[A-Za-zА-Яа-яЁё0-9._%+\-]+@[A-Za-zА-Яа-я0-9.\-]+\.[A-Za-zА-Яа-я]{2,}")
KARTA_RE = re.compile(r"(?<!\d)(?:\d{4}[\s\-]\d{4}[\s\-]\d{4}[\s\-]\d{4}|\d{16})(?!\d)")
SCHET_RE = re.compile(r"(?<!\d)\d{20}(?!\d)")
SNILS_RE = re.compile(r"(?<!\d)\d{3}[\s\-]\d{3}[\s\-]\d{3}[\s\-]?\d{2}(?!\d)")
OGRN_RE = re.compile(r"(?<!\d)(?:\d{15}|\d{13})(?!\d)")
PASPORT_RE = re.compile(r"(?<!\d)\d{4}\s\d{6}(?!\d)")
TELEFON_RE = re.compile(r"(?<!\d)(?:\+7|8|7)[\s\-]?\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2}(?!\d)")
INN_RE = re.compile(r"(?<!\d)(?:\d{12}|\d{10})(?!\d)")

# порядок важен: сначала самое узнаваемое, ИНН - последним, он самый общий
REGEX_PORYADOK = [
    ("EMAIL", EMAIL_RE),
    ("КАРТА", KARTA_RE),
    ("СЧЁТ", SCHET_RE),
    ("СНИЛС", SNILS_RE),
    ("ОГРН", OGRN_RE),
    ("ПАСПОРТ", PASPORT_RE),
    ("ТЕЛЕФОН", TELEFON_RE),
    ("ИНН", INN_RE),
]
REGEX_KATEGORII = {imya for imya, _ in REGEX_PORYADOK}

BUKVY = "АБВГДЕЖЗИКЛМНОПРСТУФХЦЧШЩЭЮЯ"


def next_metka(kategoria, schetchiki):
    """Придумывает новую метку для категории и запоминает счётчик."""
    schetchiki[kategoria] = schetchiki.get(kategoria, 0) + 1
    nomer = schetchiki[kategoria]
    if kategoria in REGEX_KATEGORII:
        return f"{kategoria}-{nomer}"
    idx = nomer - 1
    if idx < len(BUKVY):
        hvost = BUKVY[idx]
    else:
        hvost = BUKVY[idx % len(BUKVY)] + str(idx // len(BUKVY) + 1)
    return f"{kategoria}-{hvost}"


def default_slovar():
    return {
        "_kak_polzovatsya": (
            "Это твой личный словарь. Здесь хранятся настоящие данные вместе "
            "с метками, которыми они заменены в файлах из папки vhod. Никогда "
            "не отправляй этот файл в чат с нейросетью и никому не пересылай "
            "его, в нём настоящие данные. Телефоны, почту, ИНН, ОГРН, карты, "
            "счета, СНИЛС и паспорта скрипт maska.py находит сам. Если нужно "
            "спрятать имя человека или название компании, впиши его в раздел "
            "imena_i_nazvaniya: znachenie - как это написано в документах, "
            "kategoria - слово заглавными буквами, например КЛИЕНТ или "
            "ПОСТАВЩИК. Дальше всё сделает скрипт."
        ),
        "imena_i_nazvaniya": [],
        "znachenie_to_metka": {},
        "metka_to_znachenie": {},
        "schetchiki": {},
        "obrabotannye_faily": {},
    }


def load_slovar():
    if not SLOVAR_PATH.exists():
        data = default_slovar()
        save_slovar(data)
        print(f"Создал новый словарь {SLOVAR_PATH.name}.")
        return data
    try:
        with open(SLOVAR_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
    except (json.JSONDecodeError, UnicodeDecodeError):
        print(f"Не получилось прочитать словарь {SLOVAR_PATH.name} - похоже, файл повреждён.")
        print("Переименуй его (например, в slovar-staryi.json) и запусти скрипт снова,")
        print("он создаст новый пустой словарь.")
        raise SystemExit(1)
    for klyuch, znachenie in default_slovar().items():
        data.setdefault(klyuch, znachenie)
    return data


def save_slovar(data):
    with open(SLOVAR_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def mask_text(text, slovar, stats):
    """Заменяет в тексте настоящие данные на метки. Возвращает новый текст."""
    schetchiki = slovar.setdefault("schetchiki", {})
    zn_to_metka = slovar.setdefault("znachenie_to_metka", {})
    metka_to_zn = slovar.setdefault("metka_to_znachenie", {})

    # сначала имена и названия компаний из словаря
    for zapis in slovar.get("imena_i_nazvaniya", []):
        znachenie = (zapis.get("znachenie") or "").strip()
        kategoria = (zapis.get("kategoria") or "ИМЯ").strip().upper()
        if not znachenie:
            continue
        kolichestvo = text.count(znachenie)
        if kolichestvo == 0:
            continue
        klyuch = f"{kategoria}:{znachenie}"
        metka = zn_to_metka.get(klyuch)
        if metka is None:
            metka = next_metka(kategoria, schetchiki)
            zn_to_metka[klyuch] = metka
            metka_to_zn[metka] = znachenie
        text = text.replace(znachenie, metka)
        stats[kategoria] = stats.get(kategoria, 0) + kolichestvo

    # потом телефоны, почта, документы, счета и карты по шаблонам
    for kategoria, pattern in REGEX_PORYADOK:
        def zamena(m, kategoria=kategoria):
            syroe = m.group(0)
            norm = syroe.lower() if kategoria == "EMAIL" else re.sub(r"\D", "", syroe)
            klyuch = f"{kategoria}:{norm}"
            metka = zn_to_metka.get(klyuch)
            if metka is None:
                metka = next_metka(kategoria, schetchiki)
                zn_to_metka[klyuch] = metka
                metka_to_zn[metka] = syroe
            stats[kategoria] = stats.get(kategoria, 0) + 1
            return metka

        text = pattern.sub(zamena, text)

    return text


def set_paragraph_text(para, novyi_text):
    for run in list(para.runs):
        run.text = ""
    if para.runs:
        para.runs[0].text = novyi_text
    else:
        para.add_run(novyi_text)


def process_docx(path, slovar, stats, otsutstvuyut):
    try:
        import docx
    except ImportError:
        otsutstvuyut.add("Word (.docx)")
        print(f"  '{path.name}': чтобы обработать файлы Word (.docx), скрипту нужна ещё одна программа, а её сейчас нет.")
        print(f"  Если не хочешь её ставить, сохрани документ как обычный текст (.txt) и положи заново в originaly.")
        print(f"  Пока файл пропущен.")
        return None

    try:
        doc = docx.Document(str(path))
    except Exception:
        print(f"  '{path.name}': не получилось открыть файл. Возможно, он повреждён")
        print(f"  или это не настоящий файл Word. Файл пропущен.")
        return None

    def obrabotat_paragrafy(paragraphs):
        for para in paragraphs:
            if para.text.strip():
                novyi = mask_text(para.text, slovar, stats)
                if novyi != para.text:
                    set_paragraph_text(para, novyi)

    obrabotat_paragrafy(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                obrabotat_paragrafy(cell.paragraphs)
    return doc


def process_xlsx(path, slovar, stats, otsutstvuyut):
    try:
        import openpyxl
    except ImportError:
        otsutstvuyut.add("Excel (.xlsx)")
        print(f"  '{path.name}': чтобы обработать файлы Excel (.xlsx), скрипту нужна ещё одна программа, а её сейчас нет.")
        print(f"  Если не хочешь её ставить, сохрани таблицу как .csv и положи заново в originaly.")
        print(f"  Пока файл пропущен.")
        return None

    try:
        wb = openpyxl.load_workbook(str(path))
    except Exception:
        print(f"  '{path.name}': не получилось открыть файл. Возможно, он повреждён")
        print(f"  или это не настоящий файл Excel. Файл пропущен.")
        return None

    for list_ in wb.worksheets:
        for row in list_.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value.strip():
                    novoe = mask_text(cell.value, slovar, stats)
                    if novoe != cell.value:
                        cell.value = novoe
    return wb


def process_pdf(path, slovar, stats, otsutstvuyut):
    try:
        import pdfplumber
    except ImportError:
        otsutstvuyut.add("PDF")
        print(f"  '{path.name}': чтобы прочитать файлы PDF, нужна дополнительная программа pdfplumber.")
        print(f"  Установи её командой: pip install pdfplumber")
        print(f"  После установки запусти скрипт заново - пока файл пропущен.")
        return None

    try:
        kuski = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                kuski.append(page.extract_text() or "")
    except Exception:
        print(f"  '{path.name}': не получилось открыть файл. Возможно, он повреждён,")
        print(f"  защищён паролем или это не настоящий PDF. Файл пропущен.")
        return None

    polnyi_text = "\n\n".join(kuski)
    return mask_text(polnyi_text, slovar, stats)


def process_file(path, slovar, stats, otsutstvuyut):
    ext = path.suffix.lower()
    raw = path.read_bytes()
    file_hash = hashlib.md5(raw).hexdigest()
    obrab = slovar.setdefault("obrabotannye_faily", {})
    prev = obrab.get(path.name)

    if ext == ".pdf":
        out_name = f"{path.stem}-maska.txt"
    else:
        out_name = f"{path.stem}-maska{ext}"
    out_path = VHOD_DIR / out_name

    if prev and prev.get("hash") == file_hash and out_path.exists():
        print(f"  '{path.name}': уже обработан и не менялся, пропускаю")
        return False

    if ext in TEXTOVYE_RASSHIRENIYA:
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text = raw.decode("cp1251")
            except Exception:
                print(f"  '{path.name}': не получилось прочитать текст файла (незнакомая кодировка). Файл пропущен.")
                return False
        novyi_text = mask_text(text, slovar, stats)
        out_path.write_text(novyi_text, encoding="utf-8")

    elif ext == ".docx":
        doc = process_docx(path, slovar, stats, otsutstvuyut)
        if doc is None:
            return False
        doc.save(str(out_path))

    elif ext == ".xlsx":
        wb = process_xlsx(path, slovar, stats, otsutstvuyut)
        if wb is None:
            return False
        wb.save(str(out_path))

    elif ext == ".pdf":
        novyi_text = process_pdf(path, slovar, stats, otsutstvuyut)
        if novyi_text is None:
            return False
        out_path.write_text(novyi_text, encoding="utf-8")
        print(f"  '{path.name}': PDF прочитан и сохранён как обычный текст, оформление не сохраняется.")

    else:
        print(f"  '{path.name}': такой тип файла пока не поддерживается ({ext}).")
        print(f"  Поддерживаются: .txt, .md, .csv, .json, .docx, .xlsx, .pdf.")
        return False

    obrab[path.name] = {"hash": file_hash, "vhod_fail": out_name}
    return True


def wait_exit():
    try:
        input("\nНажми Enter, чтобы закрыть окно...")
    except Exception:
        pass


def main():
    try:
        ORIGINALY_DIR.mkdir(exist_ok=True)
        VHOD_DIR.mkdir(exist_ok=True)

        slovar = load_slovar()

        files = sorted(
            p for p in ORIGINALY_DIR.iterdir()
            if p.is_file() and not p.name.startswith(".")
        )

        if not files:
            print(f"В папке originaly нет файлов.")
            print(f"Положи туда файлы с настоящими данными и запусти скрипт снова.")
            wait_exit()
            return

        print(f"Нашёл файлов в originaly: {len(files)}")
        print("Начинаю обезличивание...\n")

        stats = {}
        otsutstvuyut = set()
        obrabotano = 0

        for path in files:
            print(f"Обрабатываю: {path.name}")
            try:
                if process_file(path, slovar, stats, otsutstvuyut):
                    obrabotano += 1
            except Exception as e:
                print(f"  Не получилось обработать файл '{path.name}': {e}")
                print(f"  Файл пропущен, остальные файлы будут обработаны дальше.")

        save_slovar(slovar)

        print("\nГотово.")
        print(f"Обработано файлов: {obrabotano} из {len(files)}")
        vsego = sum(stats.values())
        print(f"Всего замен: {vsego}")
        if stats:
            print("По типам:")
            for kategoria, kolichestvo in sorted(stats.items()):
                print(f"  {kategoria}: {kolichestvo}")

        if otsutstvuyut:
            spisok = ", ".join(sorted(otsutstvuyut))
            print(f"\nВнимание: часть файлов пропущена, потому что для них не хватает дополнительных программ ({spisok}).")
            print(f"Проверь папку vhod - там уже лежат обезличенные копии тех файлов, которые удалось обработать.")

        if not slovar.get("imena_i_nazvaniya"):
            print("\n" + "!" * 60)
            print("ИМЕНА И НАЗВАНИЯ КОМПАНИЙ Я НЕ ТРОГАЛ.")
            print("Телефоны, почту, ИНН и счета я нахожу сам, а имена не могу:")
            print("их не отличить от обычных слов.")
            print("")
            print("Если в твоих файлах есть фамилии клиентов или названия компаний,")
            print("которые нельзя показывать, открой файл slovar.json,")
            print("найди раздел imena_i_nazvaniya и впиши их туда.")
            print("Один раз, дальше я буду прятать их всегда.")
            print("!" * 60)

        print(f"\nОбезличенные копии лежат в папке vhod, их можно показывать нейросети.")
        print(f"ВАЖНО: файл {SLOVAR_PATH.name} хранит настоящие данные.")
        print(f"Никогда не загружай его в чат с нейросетью и никому не пересылай.")

    except Exception as e:
        print(f"Что-то пошло не так: {e}")
        print("Проверь, что рядом со скриптом есть папки originaly и vhod, и попробуй ещё раз.")

    wait_exit()


if __name__ == "__main__":
    main()
